import json
import re
from pathlib import Path
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors


ILLEGAL_CHARS_RE = re.compile(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]")


def clean_text(val):
    if isinstance(val, str):
        return ILLEGAL_CHARS_RE.sub("", val)
    return str(val)


def extract_json_blocks(md_text):
    blocks = []
    for section in md_text.split("<details>")[1:]:
        if "```json" not in section:
            continue
        try:
            json_text = section.split("```json", 1)[1].split("```", 1)[0]
            blocks.append(json.loads(json_text.strip()))
        except Exception:
            continue
    return blocks


def flatten_json(obj, parent_key="", sep="."):
    items = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            items.extend(flatten_json(v, new_key, sep))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            new_key = f"{parent_key}[{i}]"
            items.extend(flatten_json(v, new_key, sep))
    else:
        items.append((parent_key, obj))
    return items


def normalize_block(block, source_index):
    flat = dict(flatten_json(block))
    rows = []

    list_keys = [k for k in flat if "[" in k]

    if not list_keys:
        flat["source_block"] = source_index
        rows.append(flat)
        return rows

    groups = {}
    for key, val in flat.items():
        if "[" in key:
            prefix = key.split("[")[0]
            groups.setdefault(prefix, {})[key] = val
        else:
            for g in groups.values():
                g[key] = val

    for g in groups.values():
        g["source_block"] = source_index
        rows.append(g)

    return rows


def build_report(INPUT_MD: str, OUTPUT_DOCX: str, OUTPUT_PDF: str):
    INPUT_MD = Path(INPUT_MD)
    OUTPUT_DOCX = Path(OUTPUT_DOCX)
    OUTPUT_PDF = Path(OUTPUT_PDF)

    md_text = INPUT_MD.read_text(encoding="utf-8", errors="ignore")
    blocks = extract_json_blocks(md_text)

    all_rows = []
    for i, block in enumerate(blocks):
        all_rows.extend(normalize_block(block, i))

    if not all_rows:
        print("⚠ No JSON blocks found inside <details> tags.")
        return

    headers = sorted({k for r in all_rows for k in r.keys()})

    table_data = [headers]
    for r in all_rows:
        table_data.append([clean_text(r.get(h, "")) for h in headers])

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)

    # ---------- DOCX ----------
    doc = Document()
    doc.add_heading("AI GenOps Universal Analysis Report", level=1)

    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for row in table_data[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.save(OUTPUT_DOCX)

    # ---------- PDF ----------
    styles = getSampleStyleSheet()
    elements = [Paragraph("AI GenOps Universal Analysis Report", styles["Title"])]

    pdf_table = Table(table_data, repeatRows=1)
    pdf_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.25, colors.black),
        ('FONT', (0,0), (-1,0), 'Helvetica-Bold')
    ]))

    elements.append(pdf_table)

    doc_pdf = SimpleDocTemplate(str(OUTPUT_PDF))
    doc_pdf.build(elements)

    print("✔ DOCX and PDF reports generated:")
    print(" -", OUTPUT_DOCX)
    print(" -", OUTPUT_PDF)